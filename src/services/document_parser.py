"""
Document Parser — supports PDF, DOCX, MD, TXT
"""
import re
import json
import uuid
from pathlib import Path
from datetime import datetime, timezone
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument

from models.schemas import NormalizedDocument, DocumentMetadata


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}


class ParseError(Exception):
    pass


class UnsupportedFormatError(Exception):
    pass


def validate_file(file_path: Path) -> None:
    if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        ext_list = ', '.join(SUPPORTED_EXTENSIONS)
        raise UnsupportedFormatError(
            f"Formato '{file_path.suffix}' não é aceito. "
            f"Formatos aceitos: {ext_list}"
        )


def parse_document(
    file_path: Path,
    workspace_id: str = "default"
) -> tuple[NormalizedDocument, DocumentMetadata]:
    """
    Parse a document and return normalized JSON + metadata.
    Raises ParseError on failure.
    """
    validate_file(file_path)
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_path, workspace_id)
    elif ext == ".docx":
        return _parse_docx(file_path, workspace_id)
    elif ext == ".md":
        return _parse_md(file_path, workspace_id)
    elif ext == ".txt":
        return _parse_txt(file_path, workspace_id)
    else:
        raise UnsupportedFormatError(f"Formato não suportado: {ext}")


def _parse_pdf(file_path: Path, workspace_id: str):
    """Extract text from PDF using pdfplumber."""
    pages = []
    total_chars = 0

    try:
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                pages.append({
                    "page_number": i + 1,
                    "text": text
                })
                total_chars += len(text)
    except Exception as e:
        raise ParseError(f"Falha ao extrair texto do PDF: {e}")

    if total_chars == 0:
        raise ParseError("PDF não contém texto extraível (pode ser escaneado).")

    doc_id = str(uuid.uuid4())
    raw_json_path = str(file_path.parent / f"{doc_id}_raw.json")

    normalized = NormalizedDocument(
        document_id=doc_id,
        source_type="pdf",
        filename=file_path.name,
        workspace_id=workspace_id,
        created_at=datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        pages=pages,
        sections=[],
        metadata={"page_count": len(pages)},
        raw_json_path=raw_json_path
    )

    metadata = DocumentMetadata(
        document_id=doc_id,
        workspace_id=workspace_id,
        source_type="pdf",
        filename=file_path.name,
        page_count=len(pages),
        char_count=total_chars,
        chunk_count=0,
        status="parsed",
        created_at=datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        chunking_strategy="recursive"
    )

    return normalized, metadata


def _parse_docx(file_path: Path, workspace_id: str):
    """Extract text from DOCX using python-docx."""
    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        raise ParseError(f"Falha ao abrir DOCX: {e}")

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    if not full_text:
        raise ParseError("DOCX não contém texto.")

    pages = [{"page_number": 1, "text": full_text}]
    total_chars = len(full_text)

    doc_id = str(uuid.uuid4())
    raw_json_path = str(file_path.parent / f"{doc_id}_raw.json")

    normalized = NormalizedDocument(
        document_id=doc_id,
        source_type="docx",
        filename=file_path.name,
        workspace_id=workspace_id,
        created_at=datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        pages=pages,
        sections=[],
        metadata={},
        raw_json_path=raw_json_path
    )

    metadata = DocumentMetadata(
        document_id=doc_id,
        workspace_id=workspace_id,
        source_type="docx",
        filename=file_path.name,
        page_count=None,
        char_count=total_chars,
        chunk_count=0,
        status="parsed",
        created_at=datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        chunking_strategy="recursive"
    )

    return normalized, metadata


def _parse_md(file_path: Path, workspace_id: str):
    """Extract text from Markdown — split by headings."""
    try:
        content = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        content = file_path.read_text(encoding="latin-1")

    if not content.strip():
        raise ParseError("Arquivo MD vazio.")

    # Split into sections by headings
    sections = []
    lines = content.split("\n")
    current_section = {"title": "Inicio", "level": 0, "text": ""}

    for line in lines:
        # Markdown heading pattern
        m = re.match(r"^(#{1,6})\s+(.+)", line)
        if m:
            if current_section["text"]:
                sections.append(current_section)
            current_section = {
                "title": m.group(2).strip(),
                "level": len(m.group(1)),
                "text": ""
            }
        else:
            current_section["text"] += line + "\n"

    if current_section["text"]:
        sections.append(current_section)

    # Merge all text for pages
    full_text = "\n".join(s.get("text", "") for s in sections)

    pages = [{"page_number": 1, "text": full_text.strip()}]
    total_chars = len(full_text)

    doc_id = str(uuid.uuid4())
    raw_json_path = str(file_path.parent / f"{doc_id}_raw.json")

    normalized = NormalizedDocument(
        document_id=doc_id,
        source_type="md",
        filename=file_path.name,
        workspace_id=workspace_id,
        created_at=datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        pages=pages,
        sections=sections,
        metadata={"section_count": len(sections)},
        raw_json_path=raw_json_path
    )

    metadata = DocumentMetadata(
        document_id=doc_id,
        workspace_id=workspace_id,
        source_type="md",
        filename=file_path.name,
        page_count=None,
        char_count=total_chars,
        chunk_count=0,
        status="parsed",
        created_at=datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        chunking_strategy="recursive"
    )

    return normalized, metadata


def _parse_txt(file_path: Path, workspace_id: str):
    """Read plain text file."""
    try:
        content = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        content = file_path.read_text(encoding="latin-1")

    if not content.strip():
        raise ParseError("Arquivo TXT vazio.")

    pages = [{"page_number": 1, "text": content}]
    total_chars = len(content)

    doc_id = str(uuid.uuid4())
    raw_json_path = str(file_path.parent / f"{doc_id}_raw.json")

    normalized = NormalizedDocument(
        document_id=doc_id,
        source_type="txt",
        filename=file_path.name,
        workspace_id=workspace_id,
        created_at=datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        pages=pages,
        sections=[],
        metadata={},
        raw_json_path=raw_json_path
    )

    metadata = DocumentMetadata(
        document_id=doc_id,
        workspace_id=workspace_id,
        source_type="txt",
        filename=file_path.name,
        page_count=None,
        char_count=total_chars,
        chunk_count=0,
        status="parsed",
        created_at=datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        chunking_strategy="recursive"
    )

    return normalized, metadata


def save_raw_json(normalized: NormalizedDocument, output_dir: Path) -> Path:
    """Save the normalized JSON to disk."""
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"{normalized.document_id}_raw.json"
    with open(path, "w", encoding="utf-8") as f:
        json.dump(normalized.model_dump(), f, ensure_ascii=False, indent=2)
    return path
